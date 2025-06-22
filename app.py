from urllib.parse import quote
from flask import Flask, request, send_file, render_template, jsonify
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from io import BytesIO
from datetime import date
import random
import os
import re
import markdown

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

app = Flask(__name__)

from flask import send_from_directory

affirmations = [
    "Action is the foundational key to all success. – Pablo Picasso",
    "Small steps lead to big results.",
    "Discipline is choosing what you want most over what you want now.",
    "You don’t need more time. You need more focus.",
    "Start before you’re ready. Progress beats perfection.",
    "The cost of procrastination is the life you could have lived.",
    "Show up. Do the work. Trust the process.",
    "Your future is created by what you do today, not tomorrow.",
    "Consistency beats intensity every time.",
    "One focused hour today beats ten distracted ones tomorrow."
]

@app.route("/")
def home():
    prefill_goal = request.args.get("goal", "")
    return render_template("weekly.html", prefill_goal=prefill_goal)

@app.route("/daily")
def daily():
    return render_template("daily.html")

@app.route('/about')
def about():
    return render_template('about.html')

@app.route("/article/clarity")
def clarity_article():
    return render_template("clarity.html")

@app.route("/article/weekly-mission-brief")
def weekly_mission_brief_article():
    return render_template("weekly-mission-brief.html")

@app.route("/article/break-down-big-goals")
def article_break_down_goals():
    return render_template("break-down.html")

@app.route('/privacy')
def privacy():
    return render_template('privacy.html')

@app.route("/robots.txt")
def robots():
    return send_from_directory("static", "robots.txt")

@app.route("/sitemap.xml")
def sitemap():
    return send_from_directory("static", "sitemap.xml")

@app.route("/generate_daily", methods=["POST"])
def generate_daily():
    data = request.json

    # 🕷️ Honeypot spam check
    if data.get("website", "").strip():
        return jsonify({"error": "Spam detected"}), 400

    goal = data.get("goal")
    export = data.get("export", False)

    prompt = f"""
You are a tactical productivity strategist.

Break the following goal into 3–5 clearly labeled objectives. Each objective must be distinct and essential to accomplishing the overall goal.

Under each objective, list 2–4 small, specific, and actionable tasks that:
- Can be done in under an hour
- Directly contribute to completing the objective
- Are practical, not vague

Tasks must be:
- Clear, specific, and under one hour to complete
- Outcome-driven: each task should produce a visible or measurable result
- Free of fluff — ban all generic tasks like “research,” “plan,” “work on,” “review,” or “brainstorm.”
- Instead, every task must include a clear action, an object, and a method (e.g., “Find 3 blog titles using Google Trends” instead of “research blog ideas”)

Use this exact format (with bold objective titles):

Objective: Build Landing Page
- Choose a landing page builder (e.g., Carrd, Webflow, etc.)
- Draft 3 headline options
- Write 5 bullet points of benefits
- Add a signup form that connects to MailerLite

Goal: {goal}
"""



    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}]
    )

    raw = response.choices[0].message.content.strip()
    sections = raw.split("\n")
    parsed = []
    current_obj = None

    for line in sections:
        line = line.strip().replace("**", "")
        if not line:
            continue
        if line.lower().startswith("objective:"):
            if current_obj:
                parsed.append(current_obj)
            current_obj = {"objective": line.split(":", 1)[1].strip(), "tasks": []}
        elif line.startswith("- ") and current_obj:
            task_text = re.sub(r"^Task [A-Z]:\s*", "", line[2:].strip())
            current_obj["tasks"].append(task_text)

    if current_obj:
        parsed.append(current_obj)

    affirmation = random.choice(affirmations)

    if not export:
        return jsonify({
            "date": date.today().strftime("%B %d, %Y"),
            "goal": goal,
            "objectives": parsed,
            "affirmation": affirmation
        })

    # Generate Word doc
    doc = Document()
    doc.add_heading("\U0001F4C5 Stratify Daily Planner", 0)

    today = date.today().strftime("%B %d, %Y")
    date_paragraph = doc.add_paragraph(today)
    date_paragraph.runs[0].font.size = Pt(9)

    doc.add_paragraph(f"\n\U0001F3AF Today's Goal:\n{goal}", style='Intense Quote')

    tight_style = doc.styles.add_style('TightList', 1)
    tight_style.font.name = 'Calibri'
    tight_style.font.size = Pt(11)
    tight_para_format = tight_style.paragraph_format
    tight_para_format.space_before = Pt(2)
    tight_para_format.space_after = Pt(2)
    tight_para_format.line_spacing = 1.0

    for item in parsed:
        doc.add_paragraph("——————", style='TightList')
        p = doc.add_paragraph(item['objective'], style='TightList')
        p.runs[0].bold = True
        for task in item["tasks"]:
            doc.add_paragraph(f"☐ {task}", style='TightList')

    doc.add_paragraph("\n\U0001F4AC Affirmation of the Day:", style='TightList')
    doc.add_paragraph(affirmation, style='TightList')
    doc.add_paragraph("\nCrush Your Goals with StratifyPlan.com – Tap In.").runs[0].font.size = Pt(9)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name="Stratify_Daily_Plan.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/generate_weekly", methods=["POST"])
def generate_weekly():
    data = request.json

    # 🕷️ Honeypot spam check — only triggers if the field is non-empty
    if data.get("website", "").strip():
        return jsonify({"error": "Spam detected"}), 400

    goal = data.get("goal")
    export = data.get("export", False)

    prompt = f"""
You are a tactical productivity strategist.

Break the following goal into a 7-day execution plan. Assign one clear, outcome-focused objective to each day (Monday through Sunday).

Each day’s objective should:
- Progress the goal meaningfully without overlapping with other days
- Be distinct, sequenced logically, and outcome-based
- Be followed by 2–4 small, specific tasks that can be completed in under an hour each

Tasks must be:
- Clear, specific, and under one hour to complete
- Outcome-driven: each task should produce a visible or measurable result
- Free of fluff — ban all generic tasks like “research,” “plan,” “work on,” “review,” or “brainstorm.”
- Instead, every task must include a clear action, an object, and a method (e.g., “Find 3 blog titles using Google Trends” instead of “research blog ideas”)

Use this exact format:

Monday – Objective: Validate Product Idea
- Write down 3 target audience assumptions
- Create a one-question Google Form
- Post it in 2 relevant online communities
- Record 3 key takeaways from responses

Goal: {goal}
"""


    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}]
    )

    breakdown = response.choices[0].message.content.strip()
    
    # Continue your parsing/export/response logic here

    weekly_plan = []
    current_day = None
    for line in breakdown.split("\n"):
        line = line.strip().replace("**", "")
        if not line:
            continue
        match = re.match(r"^(Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday)\s*–\s*Objective:\s*(.+)", line)
        if match:
            if current_day:
                weekly_plan.append(current_day)
            current_day = {"day": match.group(1), "objective": match.group(2), "tasks": []}
        elif line.startswith("- ") and current_day:
            task_text = re.sub(r"^Task [A-Z]:\s*", "", line[2:].strip())
            current_day["tasks"].append(task_text)
    if current_day:
        weekly_plan.append(current_day)

    affirmation = random.choice(affirmations)

    if not export:
        return jsonify({
            "date": date.today().strftime("%B %d, %Y"),
            "goal": goal,
            "weekly_plan": weekly_plan,
            "affirmation": affirmation
        })

    # Generate Word doc
    doc = Document()
    doc.add_heading("\U0001F9E0 Stratify Weekly Planner", 0)

    today = date.today().strftime("%B %d, %Y")
    date_paragraph = doc.add_paragraph(today)
    date_paragraph.runs[0].font.size = Pt(9)

    doc.add_paragraph(f"\n\U0001F3AF Weekly Goal:\n{goal}", style='Intense Quote')

    tight_style = doc.styles.add_style('TightList', 1)
    tight_style.font.name = 'Calibri'
    tight_style.font.size = Pt(11)
    tight_para_format = tight_style.paragraph_format
    tight_para_format.space_before = Pt(2)
    tight_para_format.space_after = Pt(2)
    tight_para_format.line_spacing = 1.0

    for item in weekly_plan:
        doc.add_paragraph("——————", style='TightList')
        p = doc.add_paragraph(f"{item['day']} – {item['objective']}", style='TightList')
        p.runs[0].bold = True
        for task in item["tasks"]:
            doc.add_paragraph(f"☐ {task}", style='TightList')

    doc.add_paragraph("\n\U0001F4AC Affirmation of the Week:", style='TightList')
    doc.add_paragraph(affirmation, style='TightList')
    doc.add_paragraph("\nBuilt by StratifyPlan.com – The AI Planner That Doesn’t Flinch.").runs[0].font.size = Pt(9)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name="Stratify_Weekly_Plan.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


from markdown import markdown
from urllib.parse import quote

@app.route("/mission", methods=["GET", "POST"])
def mission():
    if request.method == "POST":
        goals = request.json.get("goals", "").strip()
        if not goals:
            return jsonify({
                "brief_html": "<p style='color:red;'>No input provided.</p>"
            })

        system_prompt = (
            "You are a strategic planner AI that transforms messy weekly goal input "
            "into clear, punchy 'Mission Briefs.'\n\n"
            "Use this format for each:\n\n"
            "**Objective Title –** (bold, creative, 1 line)\n"
            "Short mission statement (1–2 lines).\n"
            "Optional third line: why this matters.\n\n"
            "Output a maximum of three briefings.\n"
            "Style: tactical, sharp, cinematic."
        )

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": goals}
            ]
        )

        markdown_output = response.choices[0].message.content.strip()

        # Extract first mission title to generate URL
        linkified = ""
        for line in markdown_output.split("\n"):
            if line.startswith("**"):
                linkified = line.strip("*– ").strip()
                break

        if not linkified:
            linkified = "My Mission"


        # Convert markdown to HTML and append link
        html_content = f'<div class="mission-output">{markdown(markdown_output)}</div>'


        return jsonify({"brief_html": html_content})

    return render_template("mission.html")


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5050))
    app.run(host="0.0.0.0", port=port)
